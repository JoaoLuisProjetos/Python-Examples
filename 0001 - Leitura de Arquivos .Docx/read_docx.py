

# Must have installed docx on yout wheel '-pip install docx'
import docx
import os
import pprint



# Consumes docx file into ::var:doc
doc = docx.Document('C:/CLEAR-CASE-VOBS/SIART-20180108/SIART/01-Requisitos/Especificacoes/SIART_Espec_UC_Adiantamento_de_Valores_Provisionados.docx')
#pprint.pprint(doc.paragraphs)



# Print all text content in doc.paragraphs
for par in doc.paragraphs:
    print(par.text)